import matplotlib.pyplot as plt
from jira import JIRA
from dateutil import parser
from dateutil.relativedelta import relativedelta  # Import the relativedelta class
from docx import Document

# Helper function to calculate the median
def calculate_median(values):
    values.sort()
    n = len(values)
    if n == 0:
        return 0  # If the list is empty, return 0 as the median
    elif n % 2 == 0:
        median = (values[n // 2] + values[n // 2 - 1]) / 2
    else:
        median = values[n // 2]
    return median

# Jira credentials and connection
jira = JIRA(server='https://[domain].atlassian.net', basic_auth=('email@address.com]', '[auth_token]'))

# JQL query for retrieving all closed Features within PMZ
jql_query = 'project in (PMZ) AND issuetype in (Story, Features) AND status in (Shipped, Closed, "Release Readiness") AND updated > -60d'

# Retrieve the list of issues matching the JQL query
feature_issues = jira.search_issues(jql_query, maxResults=None)

# Get the count of features retrieved by the JQL query
feature_count = len(feature_issues)

# Define the status categories and initialize their durations
status_categories = [
    'Planning (In Progress)',
    'Waiting For Resourcing',
    'Engineering Planning',
    'Engineering In Progress',
    'Release Readiness',
]
status_durations = {status: [] for status in status_categories}

# Create a list to store the issue data
issue_data = []

# Calculate the time spent in each status for each Feature
for issue in feature_issues:
    issue = jira.issue(issue.key, expand='changelog')
    status_history = issue.changelog.histories
    last_date = parser.parse(issue.fields.created)

    for history in status_history:
        for item in history.items:
            if item.field == 'status':
                created_date = parser.parse(history.created)
                from_status = item.fromString
                to_status = item.toString

                if from_status in status_durations:
                    # Calculate the duration and add to the list
                    duration_days = (created_date - last_date).days

                    if duration_days > 0:
                        status_durations[from_status].append(duration_days)

                last_date = created_date

# Calculate the median and range (min and max) time spent in each status
median_durations = [calculate_median(status_durations[status]) for status in status_categories]
min_durations = [min(status_durations[status]) if status_durations[status] else 0 for status in status_categories]
max_durations = [max(status_durations[status]) if status_durations[status] else 0 for status in status_categories]

# Create the bar chart with error bars
plt.figure(figsize=(10, 6))
plt.bar(status_categories, median_durations)

# Add error bars representing the range from min to max duration for each status
plt.errorbar(status_categories, median_durations, yerr=[[median_durations[i] - min_durations[i] for i in range(len(status_categories))], [max_durations[i] - median_durations[i] for i in range(len(status_categories))]], fmt='none', capsize=5)

# Add labels and title
plt.xlabel('Status')
plt.ylabel('Median Time Spent (Days)')
plt.title(f'Median Time Spent in Each Status for Features in PMZ Project\nTotal Features: {feature_count}')

# Rotate the X-axis labels vertically
plt.xticks(rotation='vertical')

# Export the plot to a Word document
plt.tight_layout()  # Adjusts the spacing to prevent labels from overlapping
plt.savefig('median_status_durations.png')
plt.close()

# Save the Word document
doc = Document()
doc.add_heading('Median Feature Status Durations', level=1)

# Add the plot to the Word document
doc.add_picture('median_status_durations.png')

# Save the Word document
doc.save('median_feature_status_durations.docx')

print("Process completed successfully!")
