from jira import JIRA
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Jira credentials and connection
jira = JIRA(server='https://[domain].atlassian.net', basic_auth=('[email@address.com]', '[auth_token]'))

# JQL query for retrieving incidents involving PagerDuty within the last thirty days
jql_query = 'project in ([projects]) AND component = pagerduty-sre AND priority = "Unbreak Now!" AND created >= 2023-04-01 AND created <= 2023-05-01'

# Retrieve issues based on the JQL query
print("Retrieving issues...")
issues = jira.search_issues(jql_query, maxResults=False)
print(f"Number of issues retrieved: {len(issues)}")

# Extract necessary information and store in a DataFrame
print("Extracting information from issues...")
data = []
for i, issue in enumerate(issues, 1):
    print(f"Processing issue {i}/{len(issues)}")
    incident_id = issue.key
    created_date = issue.fields.created
    updated_date = issue.fields.updated
    resolution_date = issue.fields.resolutiondate

    created_datetime = pd.to_datetime(created_date)
    updated_datetime = pd.to_datetime(updated_date)

    # Check if resolution_datetime is not None
    if resolution_date is not None:
        resolution_datetime = pd.to_datetime(resolution_date)
        time_to_close = (resolution_datetime - created_datetime).total_seconds() / 3600
    else:
        time_to_close = None

    time_to_first_update = (updated_datetime - created_datetime).total_seconds() / 3600

    data.append([incident_id, created_datetime, time_to_first_update, time_to_close])

df = pd.DataFrame(data, columns=['Incident ID', 'Created Date', 'Time to First Update (hours)', 'Time to Close (hours)'])

# Calculate average time to update and average time to close
average_time_to_update = df['Time to First Update (hours)'].mean()
average_time_to_close = df['Time to Close (hours)'].mean()

# Create a time series chart for Time to First Update
print("Creating time series chart for Time to First Update...")
df.set_index('Created Date', inplace=True)
df['Time to First Update (hours)'] = pd.to_numeric(df['Time to First Update (hours)'], errors='coerce')
df.plot(y='Time to First Update (hours)', figsize=(10, 6))
plt.title('Time to First Update for Incidents involving PagerDuty')
plt.xlabel('Date')
plt.ylabel('Time to First Update (hours)')
plt.legend()
plt.savefig('time_to_first_update_chart.png')

# Create a time series chart for Time to Close
print("Creating time series chart for Time to Close...")
df['Time to Close (hours)'] = pd.to_numeric(df['Time to Close (hours)'], errors='coerce')
df.plot(y='Time to Close (hours)', figsize=(10, 6))
plt.title('Time to Close for Incidents involving PagerDuty')
plt.xlabel('Date')
plt.ylabel('Time to Close (hours)')
plt.legend()
plt.savefig('time_to_close_chart.png')

# Export the DataFrame to a CSV file
print("Exporting data to CSV...")
df.to_csv('incidents.csv')

# Create a Word document and insert the charts
print("Creating Word document and inserting charts...")
doc = Document()
doc.add_paragraph('Incidents involving PagerDuty')

doc.add_paragraph(f'Number of Issues: {len(issues)}')

doc.add_picture('time_to_first_update_chart.png', width=Inches(6))
doc.add_paragraph('\n')
doc.add_picture('time_to_close_chart.png', width=Inches(6))

# Add average metrics to the Word document
doc.add_paragraph(f'Average Time to Update: {average_time_to_update:.2f} hours')
doc.add_paragraph(f'Average Time to Close: {average_time_to_close:.2f} hours')

doc.save('incidents.docx')

print("Process completed successfully!")
