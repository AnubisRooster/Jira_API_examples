import matplotlib.pyplot as plt
from jira import JIRA
from dateutil import parser
from dateutil.relativedelta import relativedelta  # Import the relativedelta class
from docx import Document

# Helper function to calculate the median
def calculate_median(values):
    values.sort()
    n = len(values)
    if n == 0:
        return 0  # If the list is empty, return 0 as the median
    elif n % 2 == 0:
        median = (values[n // 2] + values[n // 2 - 1]) / 2
    else:
        median = values[n // 2]
    return median

# Jira credentials and connection
jira = JIRA(server='[jira_domain]', basic_auth=('[emailaddress]', '[auth_token]'))

# JQL query for retrieving all closed Features within PMZ
jql_query = 'project in (PMZ) AND issuetype in (Story, Features) AND updated > -360d'

# Retrieve the list of issues matching the JQL query
feature_issues = jira.search_issues(jql_query, maxResults=None)

# Define the status categories and initialize their durations
status_categories = [
    'Planning (In Progress)',
    'Waiting For Resourcing',
    'Engineering Planning',
    'Engineering In Progress',
    'Release Readiness',
]
components = [
    "component_names"
]
status_durations_by_component = {component: {status: [] for status in status_categories} for component in components}

# Calculate the time spent in each status for each Feature
for issue in feature_issues:
    issue = jira.issue(issue.key, expand='changelog')
    status_history = issue.changelog.histories
    last_date = parser.parse(issue.fields.created)

    component = "Unknown"  # Default component in case no components are assigned
    if hasattr(issue.fields, 'components') and issue.fields.components:
        component = issue.fields.components[0].name

    for history in status_history:
        for item in history.items:
            if item.field == 'status':
                created_date = parser.parse(history.created)
                from_status = item.fromString
                to_status = item.toString

                if from_status in status_categories:
                    if component in status_durations_by_component and from_status in status_durations_by_component[component]:
                        duration_days = (created_date - last_date).days

                        if duration_days > 0:
                            status_durations_by_component[component][from_status].append(duration_days)

                    last_date = created_date

# Create a Word document
doc = Document()

# Loop through each component
for component_name, component_data in status_durations_by_component.items():
    doc.add_heading(f'Median Feature Status Durations for Component {component_name}', level=1)

    median_durations = []
    min_durations = []
    max_durations = []

    # Calculate the median and range (min and max) time spent in each status for the component
    for status in status_categories:
        durations = component_data[status]
        median_durations.append(calculate_median(durations))
        min_durations.append(min(durations) if durations else 0)
        max_durations.append(max(durations) if durations else 0)

    # Add the plot to the Word document
    plt.figure(figsize=(10, 6))
    plt.bar(status_categories, median_durations)
    plt.errorbar(status_categories, median_durations, yerr=[[median_durations[i] - min_durations[i] for i in range(len(status_categories))], [max_durations[i] - median_durations[i] for i in range(len(status_categories))]], fmt='none', capsize=5)
    plt.xlabel('Status')
    plt.ylabel('Median Time Spent (Days)')
    plt.title(f'Median Time Spent in Each Status for Features in PMZ Project - Component: {component_name}')
    plt.xticks(rotation='vertical')
    plt.tight_layout()
    plt.savefig(f'median_status_durations_{component_name}.png')
    plt.close()

    doc.add_picture(f'median_status_durations_{component_name}.png')
    doc.add_page_break()

# Save the Word document
doc.save('component_median_feature_status_durations.docx')

print("Process completed successfully!")
