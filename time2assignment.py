from jira import JIRA
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Jira credentials and connection
jira = JIRA(server='https://[domain].atlassian.net', basic_auth=('[email@address.com]', '[auth_token]'))

# JQL query for retrieving incidents involving PagerDuty within the last thirty days
jql_query = 'project in (MCDB, PLAT) AND ((summary ~ "HeliosClusterPartitionUnavailable" OR summary ~ "HeliosClusterProbeFailed") AND component = pagerduty-sre AND created >= "2023-04-01" AND created <= "2023-05-01" AND priority = "Unbreak Now!")'

# Retrieve issues based on the JQL query
issues = jira.search_issues(jql_query, maxResults=False, expand='changelog')
print(f"Number of issues retrieved: {len(issues)}")

# Extract necessary information and store in a DataFrame
data = []
for i, issue in enumerate(issues, 1):
    print(f"Processing issue {i}/{len(issues)}")
    incident_id = issue.key
    created_datetime = pd.to_datetime(issue.fields.created)

    # Check if the issue has been assigned
    assigned_datetime = None
    for history in issue.changelog.histories:
        for item in history.items:
            if item.field == 'assignee':
                assigned_datetime = pd.to_datetime(history.created)
                break

    if assigned_datetime:
        time_to_assignment = (assigned_datetime - created_datetime).total_seconds() / 3600
    else:
        time_to_assignment = None

    # Check if the issue is closed
    if issue.fields.resolutiondate:
        closed_datetime = pd.to_datetime(issue.fields.resolutiondate)
        time_to_close = (closed_datetime - created_datetime).total_seconds() / 3600
    else:
        time_to_close = None

    data.append([incident_id, created_datetime, time_to_assignment, time_to_close])

df = pd.DataFrame(data, columns=['Incident ID', 'Created Date', 'Time to Assignment (hours)', 'Time to Close (hours)'])

# Calculate average time to assignment and average time to close
average_time_to_assignment = df['Time to Assignment (hours)'].mean()
average_time_to_close = df['Time to Close (hours)'].mean()

# Create a time series chart for Time to Assignment
df.set_index('Created Date', inplace=True)
df['Time to Assignment (hours)'] = pd.to_numeric(df['Time to Assignment (hours)'], errors='coerce')
df.plot(y='Time to Assignment (hours)', figsize=(10, 6))
plt.title('Time to Assignment for Incidents involving PagerDuty')
plt.xlabel('Date')
plt.ylabel('Time to Assignment (hours)')
plt.legend()
plt.savefig('time_to_assignment_chart.png')

# Create a time series chart for Time to Close
df['Time to Close (hours)'] = pd.to_numeric(df['Time to Close (hours)'], errors='coerce')
df.plot(y='Time to Close (hours)', figsize=(10, 6))
plt.title('Time to Close for Incidents involving PagerDuty')
plt.xlabel('Date')
plt.ylabel('Time to Close (hours)')
plt.legend()
plt.savefig('time_to_close_chart.png')

# Export the DataFrame to a CSV file
df.to_csv('incidents.csv')

# Create a Word document and insert the charts
doc = Document()
doc.add_paragraph('Incidents involving PagerDuty')

doc.add_paragraph(f'Number of Issues: {len(issues)}')

doc.add_picture('time_to_assignment_chart.png', width=Inches(6))
doc.add_paragraph('\n')
doc.add_picture('time_to_close_chart.png', width=Inches(6))

# Add average metrics to the Word document
doc.add_paragraph(f'Average Time to Assignment: {average_time_to_assignment:.2f} hours')
doc.add_paragraph(f'Average Time to Close: {average_time_to_close:.2f} hours')

doc.save('incidents.docx')

print("Process completed successfully!")
