import pandas as pd
from jira import JIRA
from pytz import timezone
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches

# Jira credentials and connection
jira = JIRA(server='[jira_domain]', basic_auth=('[emailaddress]', '[auth_token]'))

def add_chart_and_metrics(project_df, project_name):
    # Convert the date format to "YYYY-MM-DD"
    project_df.loc[:, 'Created Date'] = project_df['Created Date'].dt.strftime('%Y-%m-%d')

    # Scatterplot for Time in Progress
    plt.figure(figsize=(6, 4))
    plt.scatter(project_df['Created Date'], project_df['Time in Progress (hours)'])
    plt.title(f'Time in Progress for {project_name} Tickets (Excluding Weekends)')
    plt.xlabel('Date')
    plt.ylabel('Time in Progress (hours)')
    plt.savefig(f'time_in_progress_chart_{project_name}.png')
    plt.close()

    doc.add_paragraph(f'{project_name} Tickets')
    doc.add_picture(f'time_in_progress_chart_{project_name}.png', width=Inches(6))
    doc.add_paragraph('\n')

    # Scatterplot for Time to Close
    plt.figure(figsize=(6, 4))
    plt.scatter(project_df['Created Date'], project_df['Time to Close (hours)'])
    plt.title(f'Time to Close for {project_name} Tickets (Excluding Weekends)')
    plt.xlabel('Date')
    plt.ylabel('Time to Close (hours)')
    plt.savefig(f'time_to_close_chart_{project_name}.png')
    plt.close()

    doc.add_picture(f'time_to_close_chart_{project_name}.png', width=Inches(6))
    doc.add_paragraph(f'Average Time in Progress ({project_name}): {project_df["Time in Progress (hours)"].mean():.2f} hours')
    doc.add_paragraph(f'Average Time to Close ({project_name}): {project_df["Time to Close (hours)"].mean():.2f} hours')
    doc.add_page_break()

# JQL query for retrieving all closed tickets within [projects]
jql_query = 'project in ([projects]) AND issuetype in (Bug, Story, Task, Sub-task, Escalation, Investigation, "Customer Support", Test) AND status in (Closed, "Code Ready") AND created > -720d'

# Retrieve issues based on the JQL query
issues = jira.search_issues(jql_query, maxResults=False, expand='changelog')
print(f"Number of issues retrieved: {len(issues)}")

# Extract necessary information and store in a DataFrame
data = []
for i, issue in enumerate(issues, 1):
    print(f"Processing issue {i}/{len(issues)}")
    incident_id = issue.key
    created_datetime = pd.to_datetime(issue.fields.created).tz_convert('UTC')

    # Check if the issue has transitioned to "In Progress"
    in_progress_datetime = None
    for history in issue.changelog.histories:
        for item in history.items:
            if item.field == 'status' and item.toString == 'In Progress':
                in_progress_datetime = pd.to_datetime(history.created).tz_convert('UTC')
                break

    if in_progress_datetime:
        if issue.fields.resolutiondate:
            closed_datetime = pd.to_datetime(issue.fields.resolutiondate).tz_convert('UTC')

            # Exclude weekends from the time in progress calculation
            start_date = in_progress_datetime
            end_date = closed_datetime
            business_days_in_progress = pd.bdate_range(start_date, end_date)
            if len(business_days_in_progress) > 0:
                time_in_progress = np.busday_count(business_days_in_progress[0].date(), business_days_in_progress[-1].date()) * 24
            else:
                time_in_progress = None
        else:
            time_in_progress = None
    else:
        time_in_progress = None

    # Check if the issue is closed
    if issue.fields.resolutiondate:
        closed_datetime = pd.to_datetime(issue.fields.resolutiondate).tz_convert('UTC')

        # Exclude weekends from the time to close calculation
        start_date = created_datetime
        end_date = closed_datetime
        business_days_to_close = pd.bdate_range(start_date, end_date)
        if len(business_days_to_close) > 0:
            time_to_close = np.busday_count(business_days_to_close[0].date(), business_days_to_close[-1].date()) * 24
        else:
            time_to_close = None
    else:
        time_to_close = None

    data.append([incident_id, created_datetime, time_in_progress, time_to_close])

df = pd.DataFrame(data, columns=['Incident ID', 'Created Date', 'Time in Progress (hours)', 'Time to Close (hours)'])

# Calculate average time in progress and average time to close
average_time_in_progress = df['Time in Progress (hours)'].mean()
average_time_to_close = df['Time to Close (hours)'].mean()

# Create a Word document
doc = Document()

# Add individual charts and metrics for each project
projects = ['projects']
for project in projects:
    project_df = df[df['Incident ID'].str.startswith(project)]
    add_chart_and_metrics(project_df, project)

# Create a combined chart for all projects
df_combined = df.copy()
add_chart_and_metrics(df_combined, 'Combined')

# Save the Word document
doc.save('time_metrics.docx')

print("Process completed successfully!")
