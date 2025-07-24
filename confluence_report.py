import os
import sys
from jira import JIRA
import matplotlib.pyplot as plt
from docx import Document

# Jira credentials
JIRA_SERVER = os.getenv("JIRA_SERVER", 'https://[domain].atlassian.net')
DRY_RUN = True if os.getenv("DRY_RUN", "false") == "true" else False

# Hardcoded credentials (for testing purposes)
JIRA_USERNAME = "[email@address.com]"
JIRA_PASSWORD = "[auth_token]"

# Set the board ID
JIRA_BOARD_ID = "[id]"

try:
    # Connect to Jira
    jira = JIRA(server=JIRA_SERVER, basic_auth=(JIRA_USERNAME, JIRA_PASSWORD))

    closed_sprints = [s for s in jira.sprints(JIRA_BOARD_ID) if s.state == 'closed']

    if not closed_sprints:
        print("No closed sprints found for the scrum board.")
        sys.exit(1)
    else:
        # Find the last closed sprint
        sprint = max(closed_sprints, key=lambda s: s.id)
        sprint_id = sprint.id
        print(f"Last closed sprint ID: {sprint_id} {sprint.name}")

    # Get the sprint issues
    sprint_issues = jira.search_issues(f'Sprint = {sprint.id}', maxResults=False)

    # Prepare the table content
    table_rows = []
    assignee_counts = {}
    parent_task_summaries = {}  # For tasks per parent issue

    for issue in sprint_issues:
        parent_summary = "N/A"
        if hasattr(issue.fields, "parent") and issue.fields.parent:
            parent_issue = jira.issue(issue.fields.parent.key)
            parent_summary = parent_issue.fields.summary

        table_row = "| {key} | {summary} | {status} | {assignee} | {parent_summary} |\n".format(
            key=issue.key,
            summary=issue.fields.summary,
            status=issue.fields.status.name,
            assignee=issue.fields.assignee.displayName if issue.fields.assignee else "Unassigned",
            parent_summary=parent_summary,
        )
        table_rows.append(table_row)
        if issue.fields.assignee:
            assignee = issue.fields.assignee.displayName
            if assignee in assignee_counts:
                assignee_counts[assignee] += 1
            else:
                assignee_counts[assignee] = 1
        else:
            if "Unassigned" in assignee_counts:
                assignee_counts["Unassigned"] += 1
            else:
                assignee_counts["Unassigned"] = 1

        if parent_summary:
            if parent_summary in parent_task_summaries:
                parent_task_summaries[parent_summary] += 1
            else:
                parent_task_summaries[parent_summary] = 1

    # Create the bar chart for tasks per parent issue
    parent_labels = list(parent_task_summaries.keys())
    parent_values = list(parent_task_summaries.values())

    plt.figure(figsize=(6, 4))
    plt.bar(parent_labels, parent_values)
    plt.xlabel('Parent Issue')
    plt.ylabel('Count')
    plt.title(f'Sprint {sprint.name} Tasks per Parent Issue')
    plt.xticks(rotation=45)
    plt.tight_layout()

    # Save the parent task bar chart as an image
    parent_chart_file = "parent_task_bar_chart.png"
    plt.savefig(parent_chart_file)
    plt.close()

    print(f"Parent task bar chart saved successfully: {parent_chart_file}")

    # Create the bar chart for issue status
    status_counts = {}
    for issue in sprint_issues:
        status = issue.fields.status.name
        if status in status_counts:
            status_counts[status] += 1
        else:
            status_counts[status] = 1

    labels = list(status_counts.keys())
    values = list(status_counts.values())

    plt.figure(figsize=(6, 4))
    plt.bar(labels, values)
    plt.xlabel('Status')
    plt.ylabel('Count')
    plt.title(f'Sprint {sprint.name} Issue Status')
    plt.xticks(rotation=45)
    plt.tight_layout()

    # Save the status bar chart as an image
    chart_file = "status_bar_chart.png"
    plt.savefig(chart_file)
    plt.close()

    print(f"Status bar chart saved successfully: {chart_file}")

    # Create the pie chart for issues by assignee
    assignee_labels = list(assignee_counts.keys())
    assignee_values = list(assignee_counts.values())

    plt.figure(figsize=(6, 4))
    plt.pie(assignee_values, labels=assignee_labels, autopct='%1.1f%%')
    plt.title(f'Sprint {sprint.name} Issues by Assignee')
    plt.tight_layout()

    # Save the assignee pie chart as an image
    assignee_chart_file = "assignee_pie_chart.png"
    plt.savefig(assignee_chart_file)
    plt.close()

    print(f"Assignee pie chart saved successfully: {assignee_chart_file}")

    # Create the Word document
    doc = Document()
    doc.add_heading(f'Sprint {sprint.name} Summary', level=2)

    # Add the table
    table = doc.add_table(rows=1, cols=5)  # Increased column count to accommodate Parent Issue
    table.style = 'Table Grid'

    table.cell(0, 0).text = 'Issue Key'
    table.cell(0, 1).text = 'Summary'
    table.cell(0, 2).text = 'Status'
    table.cell(0, 3).text = 'Assignee'
    table.cell(0, 4).text = 'Parent Issue'  # Added Parent Issue column

    for table_row in table_rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text, cells[3].text, cells[4].text = table_row.strip().split('|')[1:-1]

    # Add the status bar chart
    doc.add_picture(chart_file)

    # Add the assignee pie chart
    doc.add_paragraph().add_run().add_picture(assignee_chart_file)

    # Add the parent task bar chart
    doc.add_paragraph().add_run().add_picture(parent_chart_file)

    # Save the Word document
    doc_file = "sprint_summary.docx"
    doc.save(doc_file)

    print(f"Document saved successfully: {doc_file}")

except Exception as e:
    print(f"Error: {e}")
